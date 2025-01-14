import os
import pandas as pd
from telegram import InlineKeyboardButton, InlineKeyboardMarkup, Update
from telegram.ext import (
    Application,
    CommandHandler,
    CallbackQueryHandler,
    ConversationHandler,
    MessageHandler,
    ContextTypes,
    filters,
)
from docx import Document
import json
import shutil
import time
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Get the bot token from .env
BOT_TOKEN = os.getenv("BOT_TOKEN")
if not BOT_TOKEN:
    raise ValueError("Bot token not found in .env file")

# Define states for the conversation
COUNTRY, USERNAME_UID, RULES = range(3)

# Links to private groups
PRIVATE_GROUPS = {
    "Ethiopia": "https://t.me/+o3zMT7hAbIU3Y2E0",
    "Nigeria": "https://t.me/+GaqELk8BynxmMmRk",
}

EXCEL_FILE_PATH = r"TournamentData.xlsx"
JSON_FILE_PATH = r"TournamentData.json"
WORD_FILE_PATH = r"TournamentData.docx"
BACKUP_FOLDER_PATH = r"backups"

# Ensure necessary files exist
def create_files_if_not_exist():
    if not os.path.exists(EXCEL_FILE_PATH):
        df = pd.DataFrame(columns=["Country", "Username", "UID", "Level"])
        with pd.ExcelWriter(EXCEL_FILE_PATH, engine="openpyxl") as writer:
            df.to_excel(writer, index=False)

    if not os.path.exists(JSON_FILE_PATH):
        with open(JSON_FILE_PATH, 'w') as json_file:
            json.dump([], json_file, indent=4)

    if not os.path.exists(WORD_FILE_PATH):
        doc = Document()
        doc.save(WORD_FILE_PATH)

    if not os.path.exists(BACKUP_FOLDER_PATH):
        os.makedirs(BACKUP_FOLDER_PATH)

# Insert player data into files
def insert_player_data_to_files(country, username, uid, level):
    new_row = {"Country": country, "Username": username, "UID": str(uid), "Level": level}

    df = pd.read_excel(EXCEL_FILE_PATH, engine="openpyxl")
    df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
    with pd.ExcelWriter(EXCEL_FILE_PATH, engine="openpyxl") as writer:
        df.to_excel(writer, index=False)

    with open(JSON_FILE_PATH, 'r') as json_file:
        data = json.load(json_file)
    data.append(new_row)
    with open(JSON_FILE_PATH, 'w') as json_file:
        json.dump(data, json_file, indent=4)

    doc = Document(WORD_FILE_PATH)
    doc.add_paragraph(f"Country: {country}")
    doc.add_paragraph(f"Username: {username}")
    doc.add_paragraph(f"UID: {uid}")
    doc.add_paragraph(f"Level: {level}")
    doc.add_paragraph("\n" + "-"*30 + "\n")
    doc.save(WORD_FILE_PATH)

# Backup files
def backup_files():
    timestamp = time.strftime("%Y%m%d%H%M%S")
    backup_folder = os.path.join(BACKUP_FOLDER_PATH, timestamp)
    os.makedirs(backup_folder)

    files_to_backup = [
        EXCEL_FILE_PATH,
        JSON_FILE_PATH,
        WORD_FILE_PATH,
    ]

    for file_path in files_to_backup:
        shutil.copy(file_path, os.path.join(backup_folder, os.path.basename(file_path)))

# Conversation handlers
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    keyboard = [
        [
            InlineKeyboardButton("Ethiopia", callback_data="Ethiopia"),
            InlineKeyboardButton("Nigeria", callback_data="Nigeria"),
        ]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text(
        "Welcome to the CODM Tournament!\nPlease select your country:", reply_markup=reply_markup
    )
    return COUNTRY

async def country_choice(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    query = update.callback_query
    await query.answer()
    context.user_data["country"] = query.data
    await query.edit_message_text(
        "You selected: {}\n\nPlease enter your username and Call of Duty UID.\n"
        "Example:\nUsername: Player123\nUID: 1234567890\nLevel:350".format(query.data)
    )
    return USERNAME_UID

async def username_uid(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    text = update.message.text
    if "Username:" in text and "UID:" in text and "Level:" in text:
        details = text.split("\n")
        username = details[0].split(":")[1].strip()
        uid = details[1].split(":")[1].strip()
        level = details[2].split(":")[1].strip()

        country = context.user_data.get("country")
        insert_player_data_to_files(country, username, uid, level)

        await update.message.reply_text(
            "Thank you for providing your details!\n\nHere are the tournament rules:\n"
            "1. Respect all participants.\n"
            "2. No cheating or using exploits.\n"
            "3. Follow the organizers' instructions.\n\nType 'agree' to accept the rules and continue."
        )
        return RULES
    else:
        await update.message.reply_text(
            "Invalid format. Please use the format:\nUsername: Player123\nUID: 1234567890\nLevel:350"
        )
        return USERNAME_UID

async def rules_agree(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    if update.message.text.lower() == "agree":
        country = context.user_data.get("country")
        private_link = PRIVATE_GROUPS.get(country, "No link available")
        await update.message.reply_text(
            f"Thank you for agreeing to the rules!\nHere is your private group link:\n{private_link}"
        )
        return ConversationHandler.END
    else:
        await update.message.reply_text("You must type 'agree' to continue.")
        return RULES

# Main function
def main():
    create_files_if_not_exist()

    application = Application.builder().token(BOT_TOKEN).build()

    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            COUNTRY: [CallbackQueryHandler(country_choice)],
            USERNAME_UID: [MessageHandler(filters.TEXT & ~filters.COMMAND, username_uid)],
            RULES: [MessageHandler(filters.TEXT & ~filters.COMMAND, rules_agree)],
        },
        fallbacks=[],
    )

    application.add_handler(conv_handler)

    print("Bot is running...")
    application.run_polling()

if __name__ == "__main__":
    main()
